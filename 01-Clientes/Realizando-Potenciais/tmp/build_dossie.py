from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(r"C:\Users\Micro\Documents\Simplifique\01-Clientes\Realizando-Potenciais")
OUT = ROOT / "Dossie_Auditoria_Realizando_Potenciais.docx"
ASSETS = ROOT / "tmp" / "dossie_assets"
ASSETS.mkdir(parents=True, exist_ok=True)

months = ["Jan", "Fev", "Mar", "Abr", "Mai", "Jun", "Jul", "Ago*"]
revenue = [3300.33, 50438.15, 58985.24, 82292.63, 23799.94, 24126.86, 42056.75, 8206.86]
costs = [42214.31, 30299.59, 23472.05, 23364.10, 33691.51, 37257.80, 35397.66, 50173.83]
results = [r-c for r,c in zip(revenue,costs)]
categories = {
    "Equipe": 115412.23,
    "Tráfego": 67906.69,
    "Contabilidade e jurídico": 56523.92,
    "Ferramentas": 18378.86,
    "Reembolsos": 14000.00,
    "Treinamentos": 1583.33,
    "Impostos e taxas": 997.81,
    "Papelaria": 590.00,
    "Registros e domínios": 478.01,
}

navy = "173B57"; blue = "2D657F"; pale = "E8F0F5"; light = "F5F8FA"; gray = "D9E1E6"
green = "147A59"; red = "B42318"; amber = "8A5A00"

def brl(x):
    s = f"{abs(x):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    return ("–" if x < 0 else "") + "R$ " + s

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tcPr.append(shd)

def margins(cell, top=90, start=100, bottom=90, end=100):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar = OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for m,v in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node=tcMar.find(qn(f"w:{m}"))
        if node is None: node=OxmlElement(f"w:{m}"); tcMar.append(node)
        node.set(qn("w:w"),str(v)); node.set(qn("w:type"),"dxa")

def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr(); rep = OxmlElement("w:tblHeader"); rep.set(qn("w:val"), "true"); trPr.append(rep)

def table(headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=str(h); shade(c,navy); margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs: run.font.bold=True; run.font.color.rgb=RGBColor(255,255,255); run.font.size=Pt(8.5)
    set_repeat_header(t.rows[0])
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,val in enumerate(row):
            c=cells[i]; c.text=str(val); margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ri%2: shade(c,light)
            for p in c.paragraphs:
                p.alignment=WD_ALIGN_PARAGRAPH.LEFT if i==0 else WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs: run.font.size=Pt(8.5); run.font.color.rgb=RGBColor(38,50,56)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return t

def heading(text, level=1):
    p=doc.add_heading(text, level=level); p.paragraph_format.keep_with_next=True
    return p

def bullet(text):
    p=doc.add_paragraph(text, style="List Bullet"); p.paragraph_format.space_after=Pt(3); return p

def pagebreak(): doc.add_page_break()

# Charts rendered with Pillow for deterministic local output.
font = ImageFont.truetype("arial.ttf", 24); small = ImageFont.truetype("arial.ttf", 18); bold = ImageFont.truetype("arialbd.ttf", 27)
def grouped_chart(path, title, labels, s1, s2, names):
    im=Image.new("RGB",(1500,620),"white"); d=ImageDraw.Draw(im); d.text((60,30),title,font=bold,fill="#111111")
    left,top,right,bottom=100,100,1440,530; d.line((left,bottom,right,bottom),fill="#555555",width=2)
    mx=max(max(s1),max(s2))*1.08; group=(right-left)/len(labels)
    for i,lbl in enumerate(labels):
        x=left+i*group+group*.18
        for j,(val,col) in enumerate(((s1[i],"#2D657F"),(s2[i],"#D38B3A"))):
            h=(val/mx)*(bottom-top); bx=x+j*group*.25; d.rectangle((bx,bottom-h,bx+group*.2,bottom),fill=col)
        d.text((x,bottom+12),lbl,font=small,fill="#333333")
    d.rectangle((1080,35,1110,55),fill="#2D657F"); d.text((1120,30),names[0],font=small,fill="#333333")
    d.rectangle((1260,35,1290,55),fill="#D38B3A"); d.text((1300,30),names[1],font=small,fill="#333333")
    im.save(path)
def result_chart(path):
    im=Image.new("RGB",(1500,560),"white"); d=ImageDraw.Draw(im); d.text((60,25),"Resultado mensal",font=bold,fill="#111111")
    left,top,right,bottom=100,90,1440,480; mx=max(abs(min(results)),max(results))*1.12; zero=top+(max(results)/mx)*(bottom-top)
    d.line((left,zero,right,zero),fill="#555555",width=2); group=(right-left)/len(months)
    for i,(lbl,val) in enumerate(zip(months,results)):
        x=left+i*group+group*.25; y=zero-(val/mx)*(bottom-top); col="#147A59" if val>=0 else "#B42318"
        d.rectangle((x,min(y,zero),x+group*.5,max(y,zero)),fill=col); d.text((x,bottom+12),lbl,font=small,fill="#333333")
    im.save(path)
def horizontal_chart(path):
    items=list(categories.items())[:6]; im=Image.new("RGB",(1450,650),"white"); d=ImageDraw.Draw(im); d.text((50,25),"Principais grupos de custo",font=bold,fill="#111111")
    mx=max(v for _,v in items); y=100
    for label,val in items:
        d.text((50,y+8),label,font=small,fill="#333333"); w=int(850*val/mx); d.rectangle((390,y,390+w,y+36),fill="#2D657F"); d.text((410+w,y+7),brl(val),font=small,fill="#333333"); y+=82
    im.save(path)
grouped_chart(ASSETS/"mensal.png","Faturamento e custos por mês",months,revenue,costs,("Faturamento","Custos"))
result_chart(ASSETS/"resultado.png")
horizontal_chart(ASSETS/"custos.png")

doc=Document(); sec=doc.sections[0]; sec.top_margin=Inches(.65); sec.bottom_margin=Inches(.65); sec.left_margin=Inches(.72); sec.right_margin=Inches(.72)
styles=doc.styles
styles["Normal"].font.name="Aptos"; styles["Normal"].font.size=Pt(10); styles["Normal"].font.color.rgb=RGBColor(38,50,56)
styles["Normal"].paragraph_format.space_after=Pt(6); styles["Normal"].paragraph_format.line_spacing=1.08
styles["Title"].font.name="Aptos Display"; styles["Title"].font.size=Pt(28); styles["Title"].font.bold=True; styles["Title"].font.color.rgb=RGBColor(0,0,0)
for n,size in ((1,18),(2,14),(3,11)):
    s=styles[f"Heading {n}"]; s.font.name="Aptos Display"; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor(0,0,0); s.paragraph_format.space_before=Pt(12); s.paragraph_format.space_after=Pt(5)

# Cover
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(70)
r=p.add_run("DOSSIÊ DE AUDITORIA"); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=RGBColor.from_string(blue)
p=doc.add_paragraph("Projeto Realizando Potenciais",style="Title"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p=doc.add_paragraph("Panorama financeiro operacional e de marketing"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].font.size=Pt(15)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(16)
r=p.add_run("Período analisado"); r.bold=True
doc.add_paragraph("Janeiro a agosto de 2026\nDados de performance e plataformas até 20 de agosto de 2026",style=None).alignment=WD_ALIGN_PARAGRAPH.CENTER
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(70); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Conclusão executiva"); r.bold=True; r.font.size=Pt(12)
p=doc.add_paragraph("O projeto está positivo no acumulado, mas a margem registrada é estreita e depende de poucos meses de forte faturamento. Três reconciliações — receita, tráfego e empréstimo POUPEX — precisam ser concluídas antes do encerramento definitivo da auditoria.")
p.alignment=WD_ALIGN_PARAGRAPH.CENTER

pagebreak(); heading("Resumo executivo")
doc.add_paragraph("Esta análise consolida a aba SCP, os extratos da Hotmart e da Principia, os investimentos de mídia informados e o relatório de redes sociais. O objetivo é separar desempenho operacional, movimentação financeira e eficiência de marketing, além de identificar lacunas que afetam a confiabilidade do resultado.")
table(["Indicador","Acumulado","Leitura"],[
    ["Faturamento SCP",brl(293206.76),"Base gerencial registrada"],
    ["Custos SCP",brl(275870.85),"94,1% do faturamento"],
    ["Resultado SCP",brl(17335.91),"Positivo, porém estreito"],
    ["Margem", "5,9%", "Sensível a oscilações mensais"],
    ["Receita líquida nas plataformas",brl(311841.60),"Antes dos saques"],
    ["Diferença de receita",brl(18634.84),"Pendente de conciliação"],
], [2.15,1.55,3.35])
heading("Leitura principal",2)
bullet("O resultado acumulado foi sustentado por fevereiro, março e abril; abril respondeu pelo maior excedente mensal.")
bullet("Janeiro, maio, junho e agosto apresentaram déficit; agosto é parcial e não deve ser comparado diretamente com meses completos.")
bullet("Equipe, tráfego e contabilidade e jurídico concentram 86,9% dos custos.")
bullet("A classificação atual inclui parcelas do POUPEX em contabilidade e jurídico, o que mistura financiamento com despesa operacional.")

pagebreak(); heading("Desempenho financeiro mensal")
table(["Mês","Faturamento","Custos","Resultado","Custos sobre faturamento"],[[months[i],brl(revenue[i]),brl(costs[i]),brl(results[i]),f"{costs[i]/revenue[i]*100:.1f}%".replace(".",",")] for i in range(8)]+[["Total",brl(sum(revenue)),brl(sum(costs)),brl(sum(results)),"94,1%"]],[.75,1.45,1.45,1.45,1.55])
doc.add_picture(str(ASSETS/"mensal.png"),width=Inches(6.9)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Abril foi o mês de melhor desempenho, com resultado de R$ 58,9 mil. Janeiro e agosto concentraram os maiores déficits. A leitura de agosto é provisória porque o período não está completo.")

pagebreak(); heading("Resultado e sustentabilidade")
doc.add_picture(str(ASSETS/"resultado.png"),width=Inches(6.9)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
heading("Interpretação",2)
doc.add_paragraph("A margem acumulada de 5,9% oferece pouca proteção contra reembolsos, aumento de mídia, oscilações de conversão ou despesas não registradas. O projeto precisa preservar um faturamento mensal acima de aproximadamente R$ 34,5 mil, média dos custos registrados, apenas para manter o equilíbrio de caixa.")
doc.add_paragraph("O padrão mensal também indica dependência de picos de venda. Sem os resultados de março e abril, o acumulado seria negativo. Isso reforça a necessidade de separar receita recorrente, lançamentos e vendas pontuais.")
heading("Cenário gerencial do financiamento",2)
doc.add_paragraph("As parcelas do POUPEX somam aproximadamente R$ 48,8 mil no período e estão classificadas como contabilidade e jurídico. Em uma visão puramente indicativa, retirando essas parcelas do custo operacional até que principal e juros sejam separados, o resultado seria próximo de R$ 66,1 mil e a margem, de 22,6%. Esse cenário não substitui a conciliação do contrato.")

pagebreak(); heading("Estrutura de custos")
table(["Categoria","Valor","Participação"],[[k,brl(v),f"{v/275870.85*100:.1f}%".replace(".",",")] for k,v in categories.items()],[3.45,1.7,1.25])
doc.add_picture(str(ASSETS/"custos.png"),width=Inches(6.8)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
heading("Pontos de atenção",2)
bullet("Equipe totalizou R$ 115,4 mil, média simples de R$ 14,4 mil por mês — abaixo da estimativa inicial de R$ 20 mil mensais.")
bullet("Tráfego somou R$ 67,9 mil na SCP, acima dos R$ 57,4 mil informados no levantamento de mídia.")
bullet("Reembolsos de R$ 14 mil devem ser reconciliados com produto, cliente e plataforma para evitar duplicidade com estornos já líquidos.")
bullet("De maio a agosto, as categorias estavam vazias na SCP; a divisão gerencial foi reconstruída pela descrição dos lançamentos e deve ser validada.")

pagebreak(); heading("Marketing e crescimento de audiência")
table(["Métrica","Resultado"],[
    ["Investimento geral informado",brl(57398.76)],
    ["Meta",brl(55427.61)],
    ["LinkedIn",brl(1100.00)],
    ["Google",brl(871.15)],
    ["Investimento em crescimento de base",brl(7986.39)],
    ["Investimento em aquisição de leads",brl(31413.51)],
],[4.6,2.0])
heading("Desempenho das redes",2)
table(["Indicador","Volume"],[["Publicações","2,38 mil"],["Impressões","5,5 milhões"],["Alcance","2,8 milhões"],["Curtidas","52,1 mil"],["Comentários","4,18 mil"],["Novos seguidores","7,84 mil"]],[4.6,2.0])
doc.add_paragraph("O Instagram concentrou aproximadamente 73% das impressões, 87% do alcance e 94% das curtidas. Maio foi o pico de impressões; julho apresentou recuperação de curtidas mesmo com menor volume de publicações.")
heading("Limite da atribuição",2)
doc.add_paragraph("Os dados disponíveis permitem medir eficiência agregada de audiência, mas não permitem calcular CPL, CAC ou ROAS por canal. Faltam quantidade de leads por origem, identificação das vendas por campanha e ligação entre UTM, cliente e receita.")

pagebreak(); heading("Achados de auditoria")
table(["Achado","Valor ou impacto","Classificação"],[
    ["Receita de plataformas acima da SCP",brl(18634.84),"Ressalva material"],
    ["Tráfego da SCP acima do valor informado",brl(10507.93),"Ressalva material"],
    ["POUPEX misturado ao custo operacional","Aprox. R$ 48,8 mil","Reclassificação necessária"],
    ["Categorias ausentes de maio a agosto","Composição inferida","Controle de cadastro"],
    ["Arquivo PrincipiaPay duplicado","9 registros","Tratado sem dupla soma"],
], [3.05,2.1,1.6])
heading("Parecer preliminar",2)
doc.add_paragraph("O projeto demonstra capacidade de geração de receita e encerra o período com resultado positivo na SCP. Contudo, o resultado de R$ 17,3 mil e a margem de 5,9% devem permanecer provisórios. A diferença entre plataformas e SCP, a divergência de tráfego e a classificação do POUPEX podem alterar materialmente a leitura final.")
heading("Riscos",2)
bullet("Risco de competência: receitas e custos podem estar registrados em meses diferentes de sua geração.")
bullet("Risco de completude: pagamentos realizados por outros CNPJs, contas ou cartões podem não estar na SCP.")
bullet("Risco de dupla contagem: reembolsos internos podem coexistir com estornos já descontados pelas plataformas.")
bullet("Risco de decisão: métricas de marketing sem atribuição podem direcionar verba para canais sem retorno comprovado.")

pagebreak(); heading("Plano de fechamento da auditoria")
table(["Prioridade","Ação","Evidência necessária","Resultado esperado"],[
    ["1","Conciliar faturamento SCP e plataformas","Extratos, saques, reservas e competência mensal","Explicar R$ 18,6 mil"],
    ["2","Conciliar mídia","Faturas Meta, Google e LinkedIn e razão SCP","Explicar R$ 10,5 mil"],
    ["3","Reclassificar POUPEX","Contrato e demonstrativo principal versus juros","Resultado operacional correto"],
    ["4","Validar equipe","Folha, prestadores, comissões e pagadores","Confirmar completude"],
    ["5","Validar reembolsos","Cliente, produto, plataforma e motivo","Eliminar duplicidades"],
    ["6","Implantar atribuição","UTM, lead, venda e canal","CPL, CAC e ROAS confiáveis"],
], [.65,2.05,2.35,1.65])
heading("Decisões recomendadas",2)
bullet("Adotar uma visão mensal única com receita por competência, caixa recebido e saldo em plataforma separados.")
bullet("Criar centros de custo específicos para equipe, mídia, ferramentas, impostos, financiamento e reembolsos.")
bullet("Estabelecer meta mínima de margem operacional e teto de custos por canal antes de ampliar investimento.")
bullet("Revisar mensalmente a conciliação entre SCP, plataformas, contas bancárias e cartões.")
heading("Conclusão",2)
doc.add_paragraph("A prioridade não é apenas reduzir custos, mas melhorar a qualidade da informação financeira. Com as três conciliações materiais concluídas, será possível determinar a margem operacional real, o retorno por canal e o nível sustentável de investimento do projeto.")

pagebreak(); heading("Fontes e critérios")
table(["Fonte","Uso na análise"],[
    ["Gastos RP atual jul 2026.xlsx — aba SCP","Faturamento, custos, equipe e classificação mensal"],
    ["Hotmart — detailed statement BRL","Movimentações, tarifas, reservas, saques e receita líquida"],
    ["Principia — financial statements","Entradas, financiamentos, parcelas, IOF e saques"],
    ["PrincipiaPay — financial statements","Verificação de duplicidade com o extrato completo"],
    ["Relatório Vanessa Redes Sociais","Publicações, alcance, impressões, engajamento e seguidores"],
    ["Informações de investimento fornecidas","Distribuição por canal e objetivo de mídia"],
], [3.15,3.55])
doc.add_paragraph("Critério financeiro: os valores da SCP representam a visão gerencial registrada. Os valores das plataformas representam geração líquida antes de saques. As duas bases não foram tratadas como equivalentes sem conciliação. Agosto é parcial.")
doc.add_paragraph("Natureza do parecer: auditoria gerencial preliminar baseada nos arquivos disponibilizados. Não constitui auditoria independente, asseguração contábil ou parecer fiscal.")

# footer
for section in doc.sections:
    p=section.footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("Dossiê de Auditoria — Projeto Realizando Potenciais"); r.font.size=Pt(8); r.font.color.rgb=RGBColor(100,110,118)

doc.save(OUT)
print(OUT)
