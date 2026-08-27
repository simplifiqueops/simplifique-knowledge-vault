from pathlib import Path
import re

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:\Users\Micro\Documents\Simplifique")
SRC = ROOT / r"01-Clientes\Realizando-Potenciais\Cúpula da decisão\Emails-T4-Transicao-Carreira-Cuidado-Humano.md"
OUT = ROOT / r"01-Clientes\Realizando-Potenciais\Cúpula da decisão\Cúpula da Decisão — 12 E-mails T4 — Revisão.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
LIGHT = "E8EEF5"


def set_font(run, name="Calibri", size=11, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, url):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(underline)
    text = OxmlElement("w:t")
    text.text = url
    run.append(rpr)
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def parse_emails(text):
    blocks = re.split(r"(?=^## E-MAIL \d{2} —)", text, flags=re.M)
    emails = []
    for block in blocks:
        if not block.startswith("## E-MAIL"):
            continue
        lines = block.strip().splitlines()
        title = lines[0][3:].strip()
        fields = {"HEADER": [], "PRÉ-HEADER": [], "CONTEÚDO": [], "LINK COM UTM": []}
        current = None
        for raw in lines[1:]:
            line = raw.strip()
            if line == "---" or not line:
                continue
            match = re.fullmatch(r"\*\*(HEADER|PRÉ-HEADER|CONTEÚDO|LINK COM UTM)\*\*", line)
            if match:
                current = match.group(1)
                continue
            if current:
                fields[current].append(line)
        emails.append((title, fields))
    return emails


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
):
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = color
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

if "Field Label" not in styles:
    field_style = styles.add_style("Field Label", WD_STYLE_TYPE.PARAGRAPH)
else:
    field_style = styles["Field Label"]
field_style.font.name = "Calibri"
field_style.font.size = Pt(9)
field_style.font.bold = True
field_style.font.color.rgb = DARK_BLUE
field_style.paragraph_format.space_before = Pt(10)
field_style.paragraph_format.space_after = Pt(3)
field_style.paragraph_format.keep_with_next = True

# Running header and footer
hp = section.header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("CÚPULA DA DECISÃO  |  MATERIAL PARA REVISÃO")
set_font(hr, size=8.5, bold=True, color=MUTED)
add_page_number(section.footer.paragraphs[0])

# Editorial cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(80)
p.paragraph_format.space_after = Pt(14)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CÚPULA DA DECISÃO")
set_font(r, size=10, bold=True, color=BLUE)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("12 E-mails T4")
set_font(r, size=29, bold=True, color=DARK_BLUE)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(24)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Transição de carreira para a área do cuidado humano")
set_font(r, size=15, color=BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(48)
r = p.add_run("VERSÃO PARA REVISÃO")
set_font(r, size=10, bold=True, color=MUTED)

table = doc.add_table(rows=3, cols=2)
table.autofit = False
table.columns[0].width = Inches(1.6)
table.columns[1].width = Inches(4.9)
meta = [
    ("Voz", "Vanessa Cesnik"),
    ("Objetivo", "Pesquisa confidencial e entrega do presente após o envio"),
    ("Público", "Pessoas em transição para a área do cuidado humano"),
]
for row, (label, value) in zip(table.rows, meta):
    row.cells[0].width = Inches(1.6)
    row.cells[1].width = Inches(4.9)
    set_cell_shading(row.cells[0], LIGHT)
    for cell in row.cells:
        set_cell_margins(cell)
    p0 = row.cells[0].paragraphs[0]
    r0 = p0.add_run(label)
    set_font(r0, size=9.5, bold=True, color=DARK_BLUE)
    p1 = row.cells[1].paragraphs[0]
    r1 = p1.add_run(value)
    set_font(r1, size=9.5)

doc.add_page_break()

emails = parse_emails(SRC.read_text(encoding="utf-8"))
for idx, (title, fields) in enumerate(emails):
    if idx:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.style = styles["Heading 1"]
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(title)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(f"REVISÃO  •  E-mail {idx + 1} de {len(emails)}")
    set_font(r, size=9, bold=True, color=MUTED)

    for label in ("HEADER", "PRÉ-HEADER"):
        p = doc.add_paragraph(label, style="Field Label")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        rr = p.add_run(" ".join(fields[label]))
        set_font(rr, size=11, bold=(label == "HEADER"), italic=(label == "PRÉ-HEADER"))

    doc.add_paragraph("CONTEÚDO", style="Field Label")
    for text in fields["CONTEÚDO"]:
        p = doc.add_paragraph()
        p.paragraph_format.widow_control = True
        r = p.add_run(text)
        set_font(r, size=11)

    doc.add_paragraph("LINK COM UTM", style="Field Label")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_hyperlink(p, " ".join(fields["LINK COM UTM"]))

props = doc.core_properties
props.title = "Cúpula da Decisão — 12 E-mails T4 — Revisão"
props.subject = "Transição de carreira para a área do cuidado humano"
props.author = "Realizando Potenciais"
props.keywords = "Cúpula da Decisão, T4, e-mails, transição de carreira, cuidado humano"

doc.save(OUT)
print(OUT)
